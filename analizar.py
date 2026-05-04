"""
Analisis de productividad de lineas de gomas MOGUL.
"""

import argparse
import datetime as dt
import re
import sys
from collections import defaultdict
from pathlib import Path

try:
    import openpyxl
except ImportError:
    sys.exit("Falta openpyxl. Instalalo con: pip install openpyxl --break-system-packages")

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("Falta python-docx. Instalalo con: pip install python-docx --break-system-packages")

COL_ORDEN = 0
COL_CAPACIDAD = 3
COL_LINEA = 5
COL_FECHA = 6
COL_PRODUCIDO = 7
COL_TIEMPO_NETO = 8
COL_NOT_CIF = 11

ETAPAS = ('formacion', 'recubrimiento', 'empaque')
ETAPA_LABEL = {
    'formacion': 'Formacion',
    'recubrimiento': 'Recubrimiento',
    'empaque': 'Empaque',
    'otro': 'Otro',
}

VENTANAS = [
    ('UAM', 365),
    ('U6M', 180),
    ('U3M', 90),
    ('U1M', 30),
    ('U7D', 7),
]


def normalizar_linea(s):
    if not s:
        return None
    return ' '.join(str(s).split())


def es_linea_gomas(linea):
    return bool(linea and 'MOGUL' in linea.upper())


def clasificar_linea(linea):
    if not linea:
        return 'otro', None
    s = ' '.join(linea.upper().split())
    m = re.search(r'MOGUL\s+([0-9A-Z]+)', s)
    tren = m.group(1) if m else None
    if s.startswith('MOGUL '):
        etapa = 'formacion'
    elif s.startswith('RECUBRIMIENTO'):
        etapa = 'recubrimiento'
    elif any(s.startswith(p) for p in ('EMBOLSADORA', 'ENVOLTURA', 'EMPAQUE')):
        etapa = 'empaque'
    else:
        etapa = 'otro'
    return etapa, tren


def orden_tren(tren):
    if tren is None:
        return (2, '')
    if tren.isdigit():
        return (0, int(tren))
    return (1, tren)


def cargar_filas(ruta_excel):
    wb = openpyxl.load_workbook(ruta_excel, data_only=True, read_only=True)
    if 'Datos' not in wb.sheetnames:
        ws = wb[wb.sheetnames[0]]
    else:
        ws = wb['Datos']
    filas = []
    max_fecha = None
    for row in ws.iter_rows(min_row=2, values_only=True):
        f = row[COL_FECHA]
        if not f:
            continue
        filas.append(row)
        if max_fecha is None or f > max_fecha:
            max_fecha = f
    return filas, max_fecha


def filtrar_mes(filas, ano, mes):
    return [r for r in filas if r[COL_FECHA].year == ano and r[COL_FECHA].month == mes]


def calcular_metricas(filas_mes, dias_corridos):
    H_CAL = dias_corridos * 24
    prod_por_linea = defaultdict(float)
    cif_min_por_linea = defaultdict(float)
    dias_con_prod = defaultdict(set)
    for r in filas_mes:
        linea = normalizar_linea(r[COL_LINEA])
        if not es_linea_gomas(linea):
            continue
        f = r[COL_FECHA].date()
        producido = r[COL_PRODUCIDO] or 0
        cif = r[COL_NOT_CIF] or 0
        prod_por_linea[linea] += producido
        cif_min_por_linea[linea] += cif
        if producido > 0:
            dias_con_prod[linea].add(f)
    resultado = {}
    for linea in sorted(prod_por_linea.keys()):
        prod = prod_por_linea[linea]
        n_dias = len(dias_con_prod[linea])
        rata = prod / n_dias if n_dias else 0
        h_prog = cif_min_por_linea[linea] / 60
        pct_util = (h_prog / H_CAL * 100) if H_CAL else 0
        etapa, tren = clasificar_linea(linea)
        resultado[linea] = {
            'etapa': etapa,
            'tren': tren,
            'producido': prod,
            'dias_con_prod': n_dias,
            'rata': rata,
            'horas_prog': h_prog,
            'horas_cal': H_CAL,
            'pct_util': pct_util,
            'alerta_cif': h_prog > H_CAL,
        }
    return resultado


def calcular_metricas_ventana(filas, fecha_inicio, fecha_fin):
    """
    Computa por linea de FORMACION en el rango [fecha_inicio, fecha_fin] (inclusivo):
      - producido total kg
      - dias con produccion (>0)
      - rata kg/dia (= producido / dias_con_prod)
      - horas calendario (= dias_corridos_en_rango * 24)
      - horas programadas (= sum(NOT_TiempoCIFMin)/60)
      - % utilizacion (= h_prog/h_cal * 100)
      - kg_esperados (= sum(CIF_min * CAPACIDAD), orden a orden, fila a fila)
      - % oee (= producido / kg_esperados * 100)
    """
    dias_rango = (fecha_fin - fecha_inicio).days + 1
    h_cal_total = dias_rango * 24

    prod_por_linea = defaultdict(float)
    cif_min_por_linea = defaultdict(float)
    kg_esperados_por_linea = defaultdict(float)
    dias_con_prod = defaultdict(set)

    for r in filas:
        f_raw = r[COL_FECHA]
        if hasattr(f_raw, 'date'):
            f_d = f_raw.date()
        else:
            f_d = f_raw
        if f_d < fecha_inicio or f_d > fecha_fin:
            continue
        linea = normalizar_linea(r[COL_LINEA])
        if not es_linea_gomas(linea):
            continue
        etapa, _ = clasificar_linea(linea)
        if etapa != 'formacion':
            continue
        producido = r[COL_PRODUCIDO] or 0
        cif_min = r[COL_NOT_CIF] or 0
        capacidad = r[COL_CAPACIDAD] or 0
        prod_por_linea[linea] += producido
        cif_min_por_linea[linea] += cif_min
        kg_esperados_por_linea[linea] += cif_min * capacidad
        if producido > 0:
            dias_con_prod[linea].add(f_d)

    resultado = {}
    for linea in prod_por_linea:
        prod = prod_por_linea[linea]
        n_dias = len(dias_con_prod[linea])
        rata = prod / n_dias if n_dias else 0
        h_prog = cif_min_por_linea[linea] / 60
        pct_util = (h_prog / h_cal_total * 100) if h_cal_total else 0
        kg_esp = kg_esperados_por_linea[linea]
        pct_oee = (prod / kg_esp * 100) if kg_esp else 0
        resultado[linea] = {
            'producido': prod,
            'dias_con_prod': n_dias,
            'rata': rata,
            'horas_prog': h_prog,
            'horas_cal': h_cal_total,
            'pct_util': pct_util,
            'kg_esperados': kg_esp,
            'pct_oee': pct_oee,
        }
    return resultado


def calcular_ventanas(filas, fecha_corte):
    """
    Devuelve dict {nombre_ventana: {linea: metricas}} para cada ventana movil.
    fecha_corte: date, fecha final inclusiva.
    """
    if hasattr(fecha_corte, 'date'):
        fecha_corte = fecha_corte.date()
    out = {}
    for nombre, dias in VENTANAS:
        fecha_inicio = fecha_corte - dt.timedelta(days=dias - 1)
        out[nombre] = calcular_metricas_ventana(filas, fecha_inicio, fecha_corte)
    return out


def lineas_formacion_ordenadas(metricas):
    """Devuelve la lista de lineas de formacion encontradas en metricas, ordenadas por tren."""
    lineas = []
    for linea, m in metricas.items():
        if m.get('etapa') == 'formacion':
            lineas.append(linea)
    def key(ln):
        _, tren = clasificar_linea(ln)
        return orden_tren(tren)
    return sorted(lineas, key=key)


def fmt_int(n):
    return f'{int(round(n)):,}'.replace(',', '.')


def fmt_dec(n, d=1):
    return f'{n:,.{d}f}'.replace(',', '_').replace('.', ',').replace('_', '.')


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def estilo_celda(cell, header=False, bold=False, align='left', size=10, color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        if align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(size)
            if header:
                run.font.bold = True
                if color is None:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    run.font.color.rgb = color
            elif bold:
                run.font.bold = True
            if color is not None and not header:
                run.font.color.rgb = color


def generar_docx(metricas, mes_label, fecha_max, dias_corridos, ventanas, ruta_salida):
    doc = Document()
    estilo_normal = doc.styles['Normal']
    estilo_normal.font.name = 'Calibri'
    estilo_normal.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    titulo = doc.add_heading('Reporte de productividad - Lineas MOGUL', level=0)
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    r = sub.add_run(f'Mes en curso: {mes_label}')
    r.font.bold = True
    r.font.size = Pt(13)
    sub.add_run(f'  -  Fecha de corte: {fecha_max.strftime("%Y-%m-%d")}  -  Dias corridos: {dias_corridos}')
    doc.add_paragraph()

    H_CAL = dias_corridos * 24
    por_etapa = defaultdict(list)
    por_tren = defaultdict(list)
    for linea, m in metricas.items():
        por_etapa[m['etapa']].append((linea, m))
        if m['tren'] is not None:
            por_tren[m['tren']].append((linea, m))

    # Resumen por etapa
    doc.add_heading('Resumen por etapa', level=1)
    intro_etapa = doc.add_paragraph()
    r = intro_etapa.add_run(
        'Las lineas trabajan en serie (formacion -> recubrimiento -> empaque) '
        'sobre el mismo producto. Por eso los kg producidos NO se suman entre '
        'etapas: cada etapa se reporta como un total independiente.'
    )
    r.font.size = Pt(10)
    r.font.italic = True

    encabezados_etapa = ['Etapa', '# lineas', 'Producido (kg)', 'H. programadas',
                         'H. calendario', '% utilizacion']
    tabla_et = doc.add_table(rows=1, cols=len(encabezados_etapa))
    tabla_et.style = 'Light Grid Accent 1'
    hdr = tabla_et.rows[0].cells
    for i, txt in enumerate(encabezados_etapa):
        hdr[i].text = txt
        set_cell_shading(hdr[i], '1F4E79')
        estilo_celda(hdr[i], header=True, align='center')

    for et in ETAPAS:
        items = por_etapa.get(et, [])
        if not items:
            continue
        n_l = len(items)
        prod_et = sum(m['producido'] for _, m in items)
        h_prog_et = sum(m['horas_prog'] for _, m in items)
        h_cal_et = H_CAL * n_l
        pct_et = (h_prog_et / h_cal_et * 100) if h_cal_et else 0
        row = tabla_et.add_row().cells
        row[0].text = ETAPA_LABEL[et]
        row[1].text = str(n_l)
        row[2].text = fmt_int(prod_et)
        row[3].text = fmt_dec(h_prog_et)
        row[4].text = fmt_int(h_cal_et)
        row[5].text = f'{fmt_dec(pct_et)} %'
        estilo_celda(row[0], bold=True)
        for i in range(1, 6):
            estilo_celda(row[i], align='right')
    doc.add_paragraph()

    # Alertas
    lineas_alerta = [(ln, m) for ln, m in metricas.items() if m['alerta_cif']]
    if lineas_alerta:
        h_alerta = doc.add_heading('Alertas de calidad de dato', level=1)
        for run in h_alerta.runs:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        intro = doc.add_paragraph()
        r = intro.add_run(
            'Las siguientes lineas reportan mas horas programadas (CIF) que horas calendario '
            'disponibles en el mes. Esto indica posible doble registro o error de captura '
            'en el archivo fuente. El reporte muestra el dato real, sin ajustes.'
        )
        r.font.size = Pt(10)
        marca = '[!]'
        for ln, m in lineas_alerta:
            exceso = m['horas_prog'] - m['horas_cal']
            p = doc.add_paragraph(style='List Bullet')
            run_b = p.add_run(ln + ': ')
            run_b.font.bold = True
            p.add_run(
                f'{fmt_dec(m["horas_prog"])} h programadas vs {fmt_int(m["horas_cal"])} h calendario '
                f'(exceso de {fmt_dec(exceso)} h, {fmt_dec(m["pct_util"])} % utilizacion).'
            )
        doc.add_paragraph()

    # Detalle por tren
    doc.add_heading('Detalle por tren', level=1)
    intro_tren = doc.add_paragraph()
    r = intro_tren.add_run(
        'Cada tren (1, 2, 3, 4, A, B) muestra sus etapas en orden: formacion -> '
        'recubrimiento -> empaque. La produccion de cada etapa es del mismo material '
        'fluyendo por el tren, no kg adicionales.'
    )
    r.font.size = Pt(10)
    r.font.italic = True

    encabezados = ['Linea', 'Etapa', 'Producido (kg)', 'Dias c/prod', 'Rata kg/dia',
                   'H. programadas', 'H. calendario', '% utilizacion']
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = 'Light Grid Accent 1'
    hdr = tabla.rows[0].cells
    for i, txt in enumerate(encabezados):
        hdr[i].text = txt
        set_cell_shading(hdr[i], '1F4E79')
        estilo_celda(hdr[i], header=True, align='center')

    color_etapa = {
        'formacion': 'E2EFDA',
        'recubrimiento': 'FFF2CC',
        'empaque': 'D9E2F3',
        'otro': 'F2F2F2',
    }
    orden_etapa = {'formacion': 0, 'recubrimiento': 1, 'empaque': 2, 'otro': 3}
    trenes_ordenados = sorted(por_tren.keys(), key=orden_tren)

    for tren in trenes_ordenados:
        sep = tabla.add_row().cells
        sep[0].text = f'Tren {tren}'
        for i in range(8):
            set_cell_shading(sep[i], 'BFBFBF')
            estilo_celda(sep[i], bold=True)
        items_tren = sorted(
            por_tren[tren],
            key=lambda x: (orden_etapa.get(x[1]['etapa'], 9), x[0])
        )
        for linea, m in items_tren:
            row = tabla.add_row().cells
            marca = ' [!]' if m['alerta_cif'] else ''
            row[0].text = linea + marca
            row[1].text = ETAPA_LABEL.get(m['etapa'], m['etapa'])
            row[2].text = fmt_int(m['producido'])
            row[3].text = str(m['dias_con_prod'])
            row[4].text = fmt_int(m['rata'])
            row[5].text = fmt_dec(m['horas_prog'])
            row[6].text = fmt_int(m['horas_cal'])
            row[7].text = f'{fmt_dec(m["pct_util"])} %'
            estilo_celda(row[0])
            estilo_celda(row[1], align='center')
            for i in range(2, 8):
                estilo_celda(row[i], align='right')
            if m['alerta_cif']:
                for i in range(8):
                    set_cell_shading(row[i], 'FCE4D6')
            else:
                set_cell_shading(row[1], color_etapa.get(m['etapa'], 'FFFFFF'))

    sin_tren = [(ln, m) for ln, m in metricas.items() if m['tren'] is None]
    if sin_tren:
        sep = tabla.add_row().cells
        sep[0].text = 'Sin tren identificado'
        for i in range(8):
            set_cell_shading(sep[i], 'BFBFBF')
            estilo_celda(sep[i], bold=True)
        for linea, m in sin_tren:
            row = tabla.add_row().cells
            row[0].text = linea
            row[1].text = ETAPA_LABEL.get(m['etapa'], m['etapa'])
            row[2].text = fmt_int(m['producido'])
            row[3].text = str(m['dias_con_prod'])
            row[4].text = fmt_int(m['rata'])
            row[5].text = fmt_dec(m['horas_prog'])
            row[6].text = fmt_int(m['horas_cal'])
            row[7].text = f'{fmt_dec(m["pct_util"])} %'
            estilo_celda(row[0])
            estilo_celda(row[1], align='center')
            for i in range(2, 8):
                estilo_celda(row[i], align='right')
    doc.add_paragraph()

    # ========== Resumen Formacion - Ventanas moviles ==========
    doc.add_heading('Resumen Formacion - Ventanas moviles', level=1)
    intro_v = doc.add_paragraph()
    r = intro_v.add_run(
        'Tablero historico de las 6 lineas de Formacion (MOGUL 1, 2, 3, 4, A, B) sobre '
        'distintos horizontes moviles: UAM (ultimo año), U6M (6 meses), U3M (3 meses), '
        'U1M (ultimo mes) y U7D (ultimos 7 dias). El OEE se calcula como '
        'kg producidos / kg esperados, donde kg esperados = sum(CIF_min × CAPACIDAD) '
        'fila por fila (orden a orden).'
    )
    r.font.size = Pt(10)
    r.font.italic = True

    # Recopilar lineas de formacion: union de todas las ventanas
    lineas_form_set = set()
    for nombre, _ in VENTANAS:
        for ln in ventanas[nombre]:
            lineas_form_set.add(ln)
    # Ordenar por tren
    def key_form(ln):
        _, tren = clasificar_linea(ln)
        return orden_tren(tren)
    lineas_form = sorted(lineas_form_set, key=key_form)

    grupos = [
        ('Kg promedio dia', 'rata', 'int'),
        ('% Utilizacion', 'pct_util', 'pct'),
        ('% OEE', 'pct_oee', 'pct'),
    ]
    n_ventanas = len(VENTANAS)
    n_cols = 1 + len(grupos) * n_ventanas  # 1 + 3*5 = 16

    tabla_v = doc.add_table(rows=2, cols=n_cols)
    tabla_v.style = 'Light Grid Accent 1'

    # Header row 1: Linea | [Kg promedio dia (5)] | [% Utilizacion (5)] | [% OEE (5)]
    fila_h1 = tabla_v.rows[0].cells
    fila_h1[0].text = 'Linea'
    set_cell_shading(fila_h1[0], '1F4E79')
    estilo_celda(fila_h1[0], header=True, align='center', size=9)
    color_grupo_hex = ['305496', '548235', 'C65911']  # azul oscuro, verde oscuro, naranja
    for gi, (g_label, _, _) in enumerate(grupos):
        c0 = 1 + gi * n_ventanas
        c1 = c0 + n_ventanas - 1
        # mergear celdas c0..c1
        merged = fila_h1[c0].merge(fila_h1[c1])
        merged.text = g_label
        set_cell_shading(merged, color_grupo_hex[gi])
        estilo_celda(merged, header=True, align='center', size=10)

    # Header row 2: '' | UAM U6M U3M U1M U7D x3
    fila_h2 = tabla_v.rows[1].cells
    fila_h2[0].text = ''
    set_cell_shading(fila_h2[0], '1F4E79')
    estilo_celda(fila_h2[0], header=True, align='center', size=9)
    for gi in range(len(grupos)):
        for vi, (vname, _) in enumerate(VENTANAS):
            ci = 1 + gi * n_ventanas + vi
            fila_h2[ci].text = vname
            set_cell_shading(fila_h2[ci], color_grupo_hex[gi])
            estilo_celda(fila_h2[ci], header=True, align='center', size=9)

    # Data rows
    for ln in lineas_form:
        row = tabla_v.add_row().cells
        # Etiqueta corta
        m = re.search(r'MOGUL\s+([0-9A-Z]+)', ln.upper())
        etiqueta = f'Mogul {m.group(1)}' if m else ln
        row[0].text = etiqueta
        estilo_celda(row[0], bold=True, align='left', size=9)
        for gi, (_, key, fmt) in enumerate(grupos):
            for vi, (vname, _) in enumerate(VENTANAS):
                ci = 1 + gi * n_ventanas + vi
                m_v = ventanas[vname].get(ln)
                if m_v is None:
                    txt = '-'
                else:
                    val = m_v.get(key, 0) or 0
                    if fmt == 'int':
                        txt = fmt_int(val)
                    elif fmt == 'pct':
                        txt = f'{fmt_dec(val)} %'
                    else:
                        txt = fmt_dec(val)
                row[ci].text = txt
                estilo_celda(row[ci], align='right', size=9)

    # Ajustar ancho aproximado de columnas (opcional)
    try:
        tabla_v.autofit = True
    except Exception:
        pass

    doc.add_paragraph()

    doc.add_heading('Notas metodologicas', level=2)
    notas = [
        'Lineas incluidas: todas las que contienen "MOGUL" en el nombre, clasificadas en formacion (MOGUL X), recubrimiento (RECUBRIMIENTO MOGUL X) y empaque (EMBOLSADORA / ENVOLTURA / EMPAQUE ... MOGUL X).',
        'Procesos en serie: las tres etapas trabajan sobre el mismo producto fisico, por lo que los kg producidos NO se suman entre etapas. Cada etapa se evalua como total independiente.',
        'Tren: identificador (1, 2, 3, 4, A, B) extraido del nombre de la linea, detras de la palabra MOGUL.',
        'Producido: suma de la columna "Producido" del archivo.',
        'Rata kg/dia (Kg promedio dia): producido total / numero de dias distintos con produccion > 0 en el horizonte. En la tabla de ventanas el horizonte es la ventana (UAM/U6M/U3M/U1M/U7D); en el detalle del mes es el mes en curso.',
        'Horas programadas: suma directa de NOT_TiempoCIFMin / 60, fila por fila.',
        'Horas calendario: dias corridos del horizonte multiplicado por 24h.',
        '% utilizacion: horas programadas / horas calendario.',
        'OEE (% OEE): kg producidos / kg esperados, donde kg esperados = sum(CIF_min × CAPACIDAD) calculado fila a fila (orden a orden) en el horizonte. CIF en minutos y CAPACIDAD tomada directamente del archivo.',
        'Ventanas moviles: UAM = 365 dias, U6M = 180 dias, U3M = 90 dias, U1M = 30 dias, U7D = 7 dias, todos terminando en la fecha de corte (inclusiva).',
        'Alerta de calidad: si para una linea las horas programadas exceden las horas calendario (en el mes en curso), se marca con [!] y se lista en la seccion "Alertas de calidad de dato". El reporte NO ajusta el calculo: revisar el archivo fuente.',
    ]
    for n in notas:
        doc.add_paragraph(n, style='List Bullet')

    doc.save(ruta_salida)
    return ruta_salida


def main():
    parser = argparse.ArgumentParser(description='Analisis productividad lineas MOGUL')
    parser.add_argument('excel', help='Ruta al archivo Excel')
    parser.add_argument('--salida', default=None, help='Ruta del .docx de salida')
    parser.add_argument('--mes', default=None, help='Mes a analizar (YYYY-MM)')
    args = parser.parse_args()

    excel_path = Path(args.excel)
    if not excel_path.exists():
        sys.exit(f'No existe el archivo: {excel_path}')

    print(f'Leyendo {excel_path} ...', flush=True)
    filas, max_fecha = cargar_filas(excel_path)
    print(f'Filas con fecha: {len(filas)}  | Fecha max: {max_fecha.date()}', flush=True)

    if args.mes:
        ano, mes = map(int, args.mes.split('-'))
        if max_fecha.year == ano and max_fecha.month == mes:
            dias_corridos = max_fecha.day
            fecha_corte = max_fecha
        else:
            if mes == 12:
                fecha_corte = dt.date(ano + 1, 1, 1) - dt.timedelta(days=1)
            else:
                fecha_corte = dt.date(ano, mes + 1, 1) - dt.timedelta(days=1)
            dias_corridos = fecha_corte.day
    else:
        ano, mes = max_fecha.year, max_fecha.month
        dias_corridos = max_fecha.day
        fecha_corte = max_fecha

    mes_label = f'{ano}-{mes:02d}'
    print(f'Analizando mes {mes_label} ({dias_corridos} dias corridos)', flush=True)

    filas_mes = filtrar_mes(filas, ano, mes)
    print(f'Filas del mes: {len(filas_mes)}', flush=True)

    metricas = calcular_metricas(filas_mes, dias_corridos)
    print(f'Lineas MOGUL encontradas: {len(metricas)}', flush=True)

    if not metricas:
        sys.exit('No se encontraron lineas MOGUL con datos en el mes solicitado.')

    if isinstance(fecha_corte, dt.datetime):
        fecha_corte_dt = fecha_corte
        fecha_corte_d = fecha_corte.date()
    else:
        fecha_corte_dt = dt.datetime.combine(fecha_corte, dt.time())
        fecha_corte_d = fecha_corte

    print('Calculando ventanas moviles (UAM/U6M/U3M/U1M/U7D)...', flush=True)
    ventanas = calcular_ventanas(filas, fecha_corte_d)
    for nombre, _ in VENTANAS:
        print(f'  {nombre}: {len(ventanas[nombre])} lineas formacion', flush=True)

    # Carpeta destino fija (Juan): OneDrive Gerencia Operaciones / Reporte.
    # Si no existe (p.ej. corriendo en Linux sandbox), caer al directorio del Excel.
    DESTINO_FIJO = Path(
        r'C:\Users\juan.alvarez\OneDrive - SUPER DE ALIMENTOS S A'
        r'\Escritorio\Gerencia de Operaciones\Reporte'
    )
    nombre_salida = f'reporte_productividad_{mes_label}.docx'

    if args.salida:
        salida = Path(args.salida)
    else:
        try:
            if DESTINO_FIJO.exists():
                salida = DESTINO_FIJO / nombre_salida
            else:
                print(
                    f'Aviso: carpeta destino {DESTINO_FIJO} no accesible, '
                    f'guardando al lado del Excel.',
                    flush=True,
                )
                salida = excel_path.parent / nombre_salida
        except OSError:
            salida = excel_path.parent / nombre_salida

    # Crear la carpeta si hace falta (solo si la ruta es valida)
    try:
        salida.parent.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        print(f'Aviso: no pude crear la carpeta {salida.parent}: {e}', flush=True)

    print(f'Generando reporte: {salida}', flush=True)
    generar_docx(metricas, mes_label, fecha_corte_dt, dias_corridos, ventanas, str(salida))
    print(f'Listo: {salida}', flush=True)


if __name__ == '__main__':
    main()
